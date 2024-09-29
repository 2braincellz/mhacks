import streamlit as st
from frontend import check_session
from frontend import get_user_id
from frontend import supabase
from docx import Document
from docx.shared import Pt
from fpdf import FPDF

if not check_session():
    st.switch_page('frontend.py')

st.write("files.py")


def create_docx(input_data, outputPath, fontSize = 12, fontName = "Cambria"):
  doc = Document()

  for line in input_data.split("\n"):
    if line.strip():
      doc.add_paragraph(line) 
    # paragraph = doc.add_paragraph(line.strip())
    # paragraph.space_after = Pt(0)  # Space after the paragraph
    # paragraph.space_before = Pt(0)  # Space before the paragraph
  for para in doc.paragraphs:
     for run in para.runs:
        font = run.font
        font.size = Pt(fontSize)
        font.name = fontName
  doc.save(outputPath)

def create_pdf(input_data, outputPath, fontSize = 12, fontName = "Times"):
  # pdf = FPDF()
  # pdf.add_page()
  # pdf.set_font(fontName, size = fontSize)
  # with open(inputPath, "r") as f:
  #   for line in f:
  #     pdf.multi_cell(210, 6, line)
  # pdf.output(outputPath)
  a4_width_mm = 210
  pt_to_mm = 0.35
  fontsize_mm = fontSize * pt_to_mm
  margin_bottom_mm = 10
  character_width_mm = 7 * pt_to_mm
  width_text = a4_width_mm / character_width_mm

  pdf = FPDF(orientation='P', unit='mm', format='A4')
  pdf.set_auto_page_break(True, margin=margin_bottom_mm)
  pdf.add_page()
  pdf.set_font(family=fontName, size=fontSize)
  

  for line in input_data.split('\n'):
    pdf.multi_cell(pdf.w - 2*pdf.l_margin, fontsize_mm, line, 0)

  pdf.output(outputPath, 'F')

def retrieve_images():

    user_id = get_user_id()
    responses = supabase.table("txt_files").select("*").eq("user_id", user_id).execute()
    
    
    # col_idx = 0

    # col1, col2, col3 = st.columns(3)

    names = [response['name'] for response in responses.data]
    data = [response['data'] for response in responses.data]

    pairs = zip(names, data)

    option = st.selectbox(
        "Choose a file",
        names
    )

    chosen_data = pairs[option]

    create_docx(chosen_data, f"{option}.docx")

    st.download_button(label = 'Download as docx', data = data, file_name = f"{option}.docx")

    create_pdf(chosen_data, f"{option}.pdf")

    st.download_button(label = 'Download as pdf', data = data, file_name = f"{option}.pdf")




        
        # if response['file']:
        #     Base64_decoded = base64.b64decode(response['file'])
        #     img_bytes = BytesIO(Base64_decoded)
        #     img = Image.open(img_bytes)

            # if col_idx % 3 == 0:
            #     with col1:
            #         st.write(img)
            # if col_idx % 3 == 1:
            #     with col2:
            #         st.write(img)
            # if col_idx % 3 == 2:
            #     with col3:
            #         st.write(img)
            # col_idx = col_idx + 1

        st.write(response)
                

def saved():


    retrieve_images()



saved()

