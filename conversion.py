from docx import Document
from docx.shared import Pt
from fpdf import FPDF

def create_docx(inputPath, outputPath, fontSize = 12, fontName = "Cambria"):
  doc = Document()

  with open(inputPath, "r") as f:
    file = f.read()
  for line in file.split("\n"):
    if line.strip():
      doc.add_paragraph(line) 
    # paragraph = doc.add_paragraph(line.strip())
    # paragraph.space_after = Pt(0)  # Space after the paragraph
    # paragraph.space_before = Pt(0)  # Space before the paragraph
  for para in doc.paragraphs:
     for run in para.runs:
        font = run.font
        font.size = Pt(fontSize)
        font.name = fontName
  doc.save(outputPath)

def create_pdf(inputPath, outputPath, fontSize = 12, fontName = "Times"):
  # pdf = FPDF()
  # pdf.add_page()
  # pdf.set_font(fontName, size = fontSize)
  # with open(inputPath, "r") as f:
  #   for line in f:
  #     pdf.multi_cell(210, 6, line)
  # pdf.output(outputPath)
  a4_width_mm = 210
  pt_to_mm = 0.35
  fontsize_mm = fontSize * pt_to_mm
  margin_bottom_mm = 10
  character_width_mm = 7 * pt_to_mm
  width_text = a4_width_mm / character_width_mm

  pdf = FPDF(orientation='P', unit='mm', format='A4')
  pdf.set_auto_page_break(True, margin=margin_bottom_mm)
  pdf.add_page()
  pdf.set_font(family=fontName, size=fontSize)
  with open(inputPath) as f:
     file = f.read()

  for line in file.split('\n'):
    pdf.multi_cell(pdf.w - 2*pdf.l_margin, fontsize_mm, line, 0)

  pdf.output(outputPath, 'F')


# create_docx("rawtext.txt", "result.docx")
# create_pdf("rawtext.txt", "result.pdf")